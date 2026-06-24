#/ =====================================================================================
#/  Admin/Director router — school management panel
#/  Manage users, classes, subjects, teacher assignments
#/  Generate logins and passwords for a class + export to DOCX
#/ =====================================================================================

from fastapi import APIRouter, Request, Depends, HTTPException, UploadFile, File, Form
from fastapi.responses import RedirectResponse, HTMLResponse, Response
from fastapi.templating import Jinja2Templates
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, delete, func
from sqlalchemy.orm import selectinload
from app.database import get_session
from app.models.models import User, Class, Subject, TeacherAssignment, Enrollment, Schedule, Grade, News, Homework
from app.auth import hash_password, require_role, require_auth
from app.logger import log_action, CREATE_USER, UPDATE_USER, DELETE_USER, CREATE_CLASS, CREATE_SUBJECT, ASSIGN_TEACHER, VIEW_LOGS, CLEANUP_LOGS
from app.main import templates
from datetime import datetime, date, timedelta
import secrets
import string
import io
import os
import aiofiles
import uuid as uuid_lib
import config

router = APIRouter(prefix="/admin", tags=["admin"])


#/ ─── Helper: generate random password / Генерация случайного пароля ────────────────
def generate_password(length=8):
    chars = string.ascii_letters + string.digits
    return ''.join(secrets.choice(chars) for _ in range(length))


#/ ─── Dashboard / Панель администратора ─────────────────────────────────────────────────
#/ Statistics, charts, recent activity summary
#/ Статистика, графики, сводка последних действий
@router.get("/dashboard", response_class=HTMLResponse)
async def admin_dashboard(
    request: Request,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(require_role("director")),
):
    #* Statistics
    students_count = (await session.execute(select(User).where(User.role == "student"))).scalars().all()
    teachers_count = (await session.execute(select(User).where(User.role == "teacher"))).scalars().all()
    classes_count = (await session.execute(select(Class))).scalars().all()
    subjects_count = (await session.execute(select(Subject))).scalars().all()

    #* Recent news (latest 3)
    recent_news = (await session.execute(
        select(News).order_by(News.created_at.desc()).limit(3)
    )).scalars().all()

    #* Recent grade entries (latest 5)
    recent_grades = (await session.execute(
        select(Grade).order_by(Grade.created_at.desc()).limit(5)
        .options(selectinload(Grade.student), selectinload(Grade.subject))
    )).scalars().all()

    #* Grade distribution (by subject for chart)
    from sqlalchemy import func
    grade_dist = (await session.execute(
        select(Grade.value, func.count(Grade.id))
        .group_by(Grade.value)
    )).all()

    grade_chart_labels = []
    grade_chart_data = []
    grade_chart_colors = []
    color_map = {5: "#22c55e", 4: "#3b82f6", 3: "#eab308", 2: "#ef4444"}
    label_map = {5: "Отлично (5)", 4: "Хорошо (4)", 3: "Удовл. (3)", 2: "Неуд. (2)"}
    for grade, count in sorted((g for g in grade_dist if g[0] is not None), key=lambda x: x[0]):
        grade_chart_labels.append(label_map.get(grade, f"Оценка {grade}"))
        grade_chart_data.append(count)
        grade_chart_colors.append(color_map.get(grade, "#6b7280"))

    #* Recent homework (latest 3)
    recent_homework = (await session.execute(
        select(Homework).order_by(Homework.created_at.desc()).limit(3)
        .options(selectinload(Homework.subject))
    )).scalars().all()

    #* Users by role
    role_counts = {}
    for r in ("student", "teacher", "director", "admin", "secretary"):
        cnt = (await session.execute(
            select(func.count(User.id)).where(User.role == r)
        )).scalar() or 0
        role_counts[r] = cnt

    #* Grade entries per day (last 7 days) for activity chart
    today_date = date.today()
    week_ago = today_date - timedelta(days=6)
    daily_rows = (await session.execute(
        select(func.date(Grade.grade_date), func.count(Grade.id))
        .where(Grade.grade_date >= week_ago)
        .group_by(func.date(Grade.grade_date))
        .order_by(func.date(Grade.grade_date))
    )).all()
    daily_map = {str(r[0]): r[1] for r in daily_rows}
    activity_labels = []
    activity_data = []
    for i in range(7):
        d = week_ago + timedelta(days=i)
        activity_labels.append(d.strftime('%d.%m'))
        activity_data.append(daily_map.get(str(d), 0))

    #* Attendance stats (this month)
    month_start = today_date.replace(day=1)
    attendance_count = (await session.execute(
        select(func.count(Grade.id))
        .where(Grade.attendance_status.isnot(None))
        .where(Grade.grade_date >= month_start)
    )).scalar() or 0

    #* Class averages (top 5 classes by avg grade)
    class_avgs_raw = (await session.execute(
        select(Class.name, func.avg(Grade.value).label('avg_val'), func.count(Grade.id).label('cnt'))
        .join(Enrollment, Enrollment.class_id == Class.id)
        .join(Grade, Grade.student_id == Enrollment.student_id)
        .where(Grade.value.isnot(None))
        .group_by(Class.id, Class.name)
        .order_by(func.avg(Grade.value).desc())
        .limit(5)
    )).all()
    class_averages = [{"name": r[0], "avg": round(r[1], 2), "count": r[2]} for r in class_avgs_raw]

    return templates.TemplateResponse("dashboard/admin.html", {
        "request": request,
        "user": user,
        "stats": {
            "students": len(students_count),
            "teachers": len(teachers_count),
            "classes": len(classes_count),
            "subjects": len(subjects_count),
            "attendance": attendance_count,
        },
        "recent_news": recent_news,
        "recent_grades": recent_grades,
        "recent_homework": recent_homework,
        "grade_chart_labels": grade_chart_labels,
        "grade_chart_data": grade_chart_data,
        "grade_chart_colors": grade_chart_colors,
        "role_counts": role_counts,
        "activity_labels": activity_labels,
        "activity_data": activity_data,
        "class_averages": class_averages,
    })


#/ ═══════════════════════════════════════════════════════════════════════════════════════
#/  USER MANAGEMENT / УПРАВЛЕНИЕ ПОЛЬЗОВАТЕЛЯМИ
#/ ═══════════════════════════════════════════════════════════════════════════════════════

#/ GET /admin/users — List all users with optional role filter
#/ Список всех пользователей с фильтром по роли
@router.get("/users", response_class=HTMLResponse)
async def list_users(
    request: Request,
    role_filter: str = "",
    session: AsyncSession = Depends(get_session),
    user: User = Depends(require_role("director")),
):
    query = select(User)
    if role_filter:
        query = query.where(User.role == role_filter)
    query = query.order_by(User.role)
    result = await session.execute(query)
    users = result.scalars().all()
    users.sort(key=lambda u: u.last_name or "")
    return templates.TemplateResponse("admin/users.html", {
        "request": request, "user": user, "users": users, "role_filter": role_filter
    })


#/ GET /admin/users/create — Select role for new user
#/ Выбор роли для нового пользователя
@router.get("/users/create", response_class=HTMLResponse)
async def create_user_role_select(
    request: Request,
    user: User = Depends(require_role("director")),
):
    return templates.TemplateResponse("admin/user_roles.html", {
        "request": request, "user": user
    })


#* Helper: save uploaded photo to disk / Сохранение загруженного фото на диск
async def save_uploaded_photo(photo: UploadFile) -> str:
    if not photo or not photo.filename:
        return ""
    ext = os.path.splitext(photo.filename)[1] or ".jpg"
    filename = f"{uuid_lib.uuid4().hex}{ext}"
    filepath = os.path.join(config.UPLOAD_DIR, filename)
    os.makedirs(config.UPLOAD_DIR, exist_ok=True)
    content = await photo.read()
    async with aiofiles.open(filepath, "wb") as f:
        await f.write(content)
    return f"/uploads/{filename}"


#/ ─── CREATE: Teacher / СОЗДАНИЕ: Учитель ──────────────────────────────────────────
@router.get("/users/create/teacher", response_class=HTMLResponse)
async def create_teacher_page(
    request: Request,
    user: User = Depends(require_role("director")),
):
    return templates.TemplateResponse("admin/create_teacher.html", {
        "request": request, "user": user
    })


@router.post("/users/create/teacher")
async def create_teacher(
    request: Request,
    session: AsyncSession = Depends(get_session),
    admin: User = Depends(require_role("director")),
):
    form = await request.form()
    photo_file = form.get("photo")
    return await _create_user_handler(request, session, admin, form, photo_file, "teacher",
                                      template="admin/create_teacher.html")


#/ ─── CREATE: Student / СОЗДАНИЕ: Ученик ──────────────────────────────────────────
@router.get("/users/create/student", response_class=HTMLResponse)
async def create_student_page(
    request: Request,
    user: User = Depends(require_role("director")),
):
    return templates.TemplateResponse("admin/create_student.html", {
        "request": request, "user": user
    })


@router.post("/users/create/student")
async def create_student(
    request: Request,
    session: AsyncSession = Depends(get_session),
    admin: User = Depends(require_role("director")),
):
    form = await request.form()
    return await _create_user_handler(request, session, admin, form, None, "student",
                                      template="admin/create_student.html")


#/ ─── CREATE: Other Staff / СОЗДАНИЕ: Другой персонал ──────────────────────────────
@router.get("/users/create/other", response_class=HTMLResponse)
async def create_other_page(
    request: Request,
    user: User = Depends(require_role("director")),
):
    return templates.TemplateResponse("admin/create_other.html", {
        "request": request, "user": user
    })


@router.post("/users/create/other")
async def create_other(
    request: Request,
    session: AsyncSession = Depends(get_session),
    admin: User = Depends(require_role("director")),
):
    form = await request.form()
    photo_file = form.get("photo")
    return await _create_user_handler(request, session, admin, form, photo_file, "other",
                                      template="admin/create_other.html")


_TRANSLIT = {
    'а': 'a', 'б': 'b', 'в': 'v', 'г': 'g', 'д': 'd', 'е': 'e', 'ё': 'e',
    'ж': 'zh', 'з': 'z', 'и': 'i', 'й': 'j', 'к': 'k', 'л': 'l', 'м': 'm',
    'н': 'n', 'о': 'o', 'п': 'p', 'р': 'r', 'с': 's', 'т': 't', 'у': 'u',
    'ф': 'f', 'х': 'h', 'ц': 'c', 'ч': 'ch', 'ш': 'sh', 'щ': 'shch', 'ъ': '',
    'ы': 'y', 'ь': '', 'э': 'e', 'ю': 'yu', 'я': 'ya',
    'А': 'A', 'Б': 'B', 'В': 'V', 'Г': 'G', 'Д': 'D', 'Е': 'E', 'Ё': 'E',
    'Ж': 'Zh', 'З': 'Z', 'И': 'I', 'Й': 'J', 'К': 'K', 'Л': 'L', 'М': 'M',
    'Н': 'N', 'О': 'O', 'П': 'P', 'Р': 'R', 'С': 'S', 'Т': 'T', 'У': 'U',
    'Ф': 'F', 'Х': 'H', 'Ц': 'C', 'Ч': 'Ch', 'Ш': 'Sh', 'Щ': 'Shch', 'Ъ': '',
    'Ы': 'Y', 'Ь': '', 'Э': 'E', 'Ю': 'Yu', 'Я': 'Ya',
}


def _slugify(text: str) -> str:
    return "".join(_TRANSLIT.get(c, c) for c in text)


#! Handler: create user of any role (teacher/student/staff). Validates, hashes password, logs.
#! Обработчик: создание пользователя любой роли. Валидация, хэширование пароля, логирование.
async def _create_user_handler(request, session, admin, form, photo_file, role, template):
    password = form.get("password", "").strip() or generate_password()
    first_name = form.get("first_name", "").strip()
    last_name = form.get("last_name", "").strip()
    middle_name = form.get("middle_name", "").strip() or None
    phone = form.get("phone", "").strip() or None



    if not all([first_name, last_name]):
        return templates.TemplateResponse(template, {
            "request": request, "user": admin,
            "error": "Заполните обязательные поля"
        }, status_code=400)

    #* For students — auto generate internal email and login (username)
    if role == "student":
        email = f"s-{uuid_lib.uuid4().hex[:8]}@l"
        base_username = _slugify(f"{last_name}.{first_name}").lower()
        #* Make unique
        username = base_username
        for i in range(1, 100):
            existing_u = await session.execute(select(User).where(User.username == username))
            if not existing_u.scalar_one_or_none():
                break
            username = f"{base_username}{i}"
    else:
        email = form.get("email", "").strip().lower()
        username = None
        if not email:
            return templates.TemplateResponse(template, {
                "request": request, "user": admin,
                "error": "Заполните email"
            }, status_code=400)

        existing = await session.execute(select(User).where(User.email == email))
        if existing.scalar_one_or_none():
            return templates.TemplateResponse(template, {
                "request": request, "user": admin,
                "error": "Email уже используется"
            }, status_code=400)

    photo_url = ""
    if photo_file and hasattr(photo_file, 'filename') and photo_file.filename:
        try:
            photo_url = await save_uploaded_photo(photo_file)
        except Exception:
            pass

    user_kwargs = dict(
        email=email, password_hash=hash_password(password), role=role,
        first_name=first_name, last_name=last_name, middle_name=middle_name,
        phone=phone, photo_url=photo_url or None,
    )
    if username:
        user_kwargs["username"] = username

    if role == "teacher":
        labor_book = form.get("labor_book_number", "").strip() or None
        exp_str = form.get("experience_years", "").strip()
        experience = int(exp_str) if exp_str and exp_str.isdigit() else None
        salary_str = form.get("salary_monthly", "").strip()
        salary = float(salary_str) if salary_str else None
        hours_str = form.get("hours_per_week", "").strip()
        hours = float(hours_str) if hours_str else None
        user_kwargs["labor_book_number"] = labor_book
        user_kwargs["experience_years"] = experience
        user_kwargs["salary_monthly"] = salary
        user_kwargs["hours_per_week"] = hours

    elif role == "student":
        dob_str = form.get("date_of_birth", "").strip()
        if dob_str:
            try:
                user_kwargs["date_of_birth"] = datetime.strptime(dob_str, "%Y-%m-%d").date()
            except ValueError:
                pass
        user_kwargs["address"] = form.get("address", "").strip() or None
        user_kwargs["mother_name"] = form.get("mother_name", "").strip() or None
        user_kwargs["mother_phone"] = form.get("mother_phone", "").strip() or None
        user_kwargs["father_name"] = form.get("father_name", "").strip() or None
        user_kwargs["father_phone"] = form.get("father_phone", "").strip() or None
        user_kwargs["emergency_contact"] = form.get("emergency_contact", "").strip() or None
        user_kwargs["medical_info"] = form.get("medical_info", "").strip() or None

    elif role == "other":
        user_kwargs["staff_position"] = form.get("staff_position", "").strip() or None

    new_user = User(**user_kwargs)
    session.add(new_user)
    await session.commit()

    log_action(user_id=str(admin.id), action=CREATE_USER,
               details={"target_user_id": str(new_user.id), "role": role, "email": email})
    return RedirectResponse(url="/admin/users", status_code=302)


#/ GET /admin/users/{user_id}/edit — Edit user form
#/ Форма редактирования пользователя
@router.get("/users/{user_id}/edit", response_class=HTMLResponse)
async def edit_user_page(
    request: Request, user_id: str,
    session: AsyncSession = Depends(get_session),
    admin: User = Depends(require_role("director")),
):
    result = await session.execute(select(User).where(User.id == user_id))
    target = result.scalar_one_or_none()
    if not target:
        raise HTTPException(status_code=404, detail="Пользователь не найден")
    return templates.TemplateResponse("admin/user_form.html", {
        "request": request, "user": admin, "edit": True, "target": target
    })


#! POST /admin/users/{user_id}/edit — Update user data + optional password change
#! Обновление данных пользователя + опциональная смена пароля
@router.post("/users/{user_id}/edit")
async def edit_user(
    request: Request, user_id: str,
    session: AsyncSession = Depends(get_session),
    admin: User = Depends(require_role("director")),
):
    result = await session.execute(select(User).where(User.id == user_id))
    target = result.scalar_one_or_none()
    if not target:
        raise HTTPException(status_code=404, detail="Пользователь не найден")

    form = await request.form()
    target.email = form.get("email", target.email).strip().lower()
    target.first_name = form.get("first_name", target.first_name).strip()
    target.last_name = form.get("last_name", target.last_name).strip()
    target.middle_name = form.get("middle_name", "").strip() or None
    target.phone = form.get("phone", "").strip() or None

    #* Personal data
    dob_str = form.get("date_of_birth", "").strip()
    if dob_str:
        try:
            target.date_of_birth = datetime.strptime(dob_str, "%Y-%m-%d").date()
        except ValueError:
            pass
    target.address = form.get("address", "").strip() or None
    target.mother_name = form.get("mother_name", "").strip() or None
    target.mother_phone = form.get("mother_phone", "").strip() or None
    target.father_name = form.get("father_name", "").strip() or None
    target.father_phone = form.get("father_phone", "").strip() or None
    target.emergency_contact = form.get("emergency_contact", "").strip() or None
    target.medical_info = form.get("medical_info", "").strip() or None
    target.is_active = form.get("is_active", "on") == "on"

    new_password = form.get("password", "").strip()
    if new_password:
        target.password_hash = hash_password(new_password)

    await session.commit()
    log_action(user_id=str(admin.id), action=UPDATE_USER, details={"target_user_id": user_id})
    return RedirectResponse(url="/admin/users", status_code=302)


#! POST /admin/users/{user_id}/delete — Delete user (cannot delete director)
#! Удаление пользователя (нельзя удалить директора)
@router.post("/users/{user_id}/delete")
async def delete_user(
    request: Request, user_id: str,
    session: AsyncSession = Depends(get_session),
    admin: User = Depends(require_role("director")),
):
    result = await session.execute(select(User).where(User.id == user_id))
    target = result.scalar_one_or_none()
    if not target:
        raise HTTPException(status_code=404, detail="Пользователь не найден")
    if target.role == "director":
        raise HTTPException(status_code=403, detail="Нельзя удалить директора")

    await session.delete(target)
    await session.commit()
    log_action(user_id=str(admin.id), action=DELETE_USER, details={"target_user_id": user_id})
    return RedirectResponse(url="/admin/users", status_code=302)


#/ ─── School performance report / Отчёт об успеваемости школы ────────────────────────
#! Shows average grades per class, subject, and overall school statistics
#! Показывает средние оценки по классам, предметам и общую статистику
@router.get("/reports/school", response_class=HTMLResponse)
async def school_report(
    request: Request,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(require_role("director")),
):
    # Get all classes
    classes_result = await session.execute(select(Class).order_by(Class.name))
    classes = classes_result.scalars().all()
    
    # Get all subjects
    subjects_result = await session.execute(select(Subject).order_by(Subject.name))
    subjects = subjects_result.scalars().all()
    
    # Per-class averages (across all subjects)
    class_stats = []
    for cls in classes:
        # Get all students in this class
        students = await session.execute(
            select(User.id).join(Enrollment).where(Enrollment.class_id == cls.id, User.role == "student")
        )
        student_ids = [r[0] for r in students.all()]
        
        if student_ids:
            # Get average grade for the class
            avg_result = await session.execute(
                select(func.avg(Grade.value)).where(
                    Grade.student_id.in_(student_ids),
                    Grade.value.isnot(None),
                    Grade.grade_type == "regular"
                )
            )
            avg = avg_result.scalar()
            
            # Count grades
            count_result = await session.execute(
                select(func.count(Grade.id)).where(
                    Grade.student_id.in_(student_ids),
                    Grade.value.isnot(None),
                    Grade.grade_type == "regular"
                )
            )
            grade_count = count_result.scalar()
            
            # Attendance stats
            absent_result = await session.execute(
                select(func.count(Grade.id)).where(
                    Grade.student_id.in_(student_ids),
                    Grade.attendance_status == "absent"
                )
            )
            absent = absent_result.scalar()
            
            sick_result = await session.execute(
                select(func.count(Grade.id)).where(
                    Grade.student_id.in_(student_ids),
                    Grade.attendance_status == "sick"
                )
            )
            sick = sick_result.scalar()
        else:
            avg = None
            grade_count = 0
            absent = 0
            sick = 0
        
        class_stats.append({
            "class": cls,
            "average": round(avg, 2) if avg else None,
            "grade_count": grade_count,
            "absent": absent,
            "sick": sick,
            "student_count": len(student_ids) if student_ids else 0,
        })
    
    # School-wide averages per subject
    subject_stats = []
    for subj in subjects:
        avg_result = await session.execute(
            select(func.avg(Grade.value)).where(
                Grade.subject_id == subj.id,
                Grade.grade_type == "regular",
                Grade.value.isnot(None),
            )
        )
        avg = avg_result.scalar()
        subject_stats.append({
            "subject": subj,
            "average": round(avg, 2) if avg else None,
        })
    
    # Overall school average
    school_avg_result = await session.execute(
        select(func.avg(Grade.value)).where(
            Grade.grade_type == "regular",
            Grade.value.isnot(None),
        )
    )
    school_avg = school_avg_result.scalar()
    
    return templates.TemplateResponse("admin/report_school.html", {
        "request": request, "user": user,
        "class_stats": class_stats,
        "subject_stats": subject_stats,
        "school_average": round(school_avg, 2) if school_avg else None,
    })


#/ ─── Class performance report / Отчёт об успеваемости класса ────────────────────────
#! Shows detailed per-subject grades for each student in a class
#! Показывает подробные оценки по предметам для каждого ученика класса
@router.get("/reports/class/{class_id}", response_class=HTMLResponse)
async def class_report(
    request: Request, class_id: str,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(require_role("director")),
):
    # Get class
    class_result = await session.execute(select(Class).where(Class.id == class_id))
    class_ = class_result.scalar_one_or_none()
    if not class_:
        raise HTTPException(status_code=404)
    
    # Get students in this class
    students_result = await session.execute(
        select(User).join(Enrollment).where(
            Enrollment.class_id == class_id, User.role == "student"
        )
    )
    students = students_result.scalars().all()
    students.sort(key=lambda u: (u.last_name or "", u.first_name or ""))
    
    # Get all subjects
    subjects_result = await session.execute(select(Subject).order_by(Subject.name))
    subjects = subjects_result.scalars().all()
    
    # For each student, get grades per subject
    student_data = []
    for student in students:
        subject_grades = []
        total = 0
        count = 0
        for subj in subjects:
            grades_result = await session.execute(
                select(Grade).where(
                    Grade.student_id == student.id,
                    Grade.subject_id == subj.id,
                    Grade.grade_type == "regular",
                    Grade.value.isnot(None),
                ).order_by(Grade.grade_date.desc())
            )
            grades = grades_result.scalars().all()
            
            if grades:
                subj_avg = sum(g.value for g in grades if g.value) / len([g for g in grades if g.value])
                total += sum(g.value for g in grades if g.value)
                count += len([g for g in grades if g.value])
            else:
                subj_avg = None
            
            subject_grades.append({
                "subject": subj,
                "average": round(subj_avg, 2) if subj_avg else None,
                "grades": grades,
                "count": len(grades),
            })
        
        student_data.append({
            "student": student,
            "subjects": subject_grades,
            "overall_avg": round(total / count, 2) if count > 0 else None,
        })
    
    # Class average per subject
    class_subject_avgs = []
    for subj in subjects:
        avg_result = await session.execute(
            select(func.avg(Grade.value)).where(
                Grade.subject_id == subj.id,
                Grade.grade_type == "regular",
                Grade.value.isnot(None),
            )
        )
        avg = avg_result.scalar()
        class_subject_avgs.append({
            "subject": subj,
            "average": round(avg, 2) if avg else None,
        })
    
    return templates.TemplateResponse("admin/report_class.html", {
        "request": request, "user": user,
        "class": class_,
        "student_data": student_data,
        "class_subject_avgs": class_subject_avgs,
        "subjects": subjects,
    })


#/ ─── Salary report / Зарплатный отчёт ──────────────────────────────────────────────
#! Shows all teachers with salary, hours, calculated hourly rate
#! Показывает всех учителей с зарплатой, часами и расчитанной ставкой в час
@router.get("/reports/salary", response_class=HTMLResponse)
async def salary_report(
    request: Request,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(require_role("admin")),
):
    #* Get all teachers with salary data
    result = await session.execute(
        select(User).where(User.role == "teacher")
    )
    teachers = result.scalars().all()
    teachers.sort(key=lambda u: (u.last_name or "", u.first_name or ""))

    teacher_data = []
    total_salary = 0
    for t in teachers:
        hourly = None
        if t.salary_monthly and t.hours_per_week and t.hours_per_week > 0:
            #* Standard calculation: monthly / (weekly_hours * 4.33 weeks per month)
            hourly = round(t.salary_monthly / (t.hours_per_week * 4.33), 2)
        total_salary += t.salary_monthly or 0
        teacher_data.append({
            "teacher": t,
            "hourly_rate": hourly,
        })

    return templates.TemplateResponse("admin/report_salary.html", {
        "request": request, "user": user,
        "teacher_data": teacher_data,
        "total_salary": total_salary,
        "teacher_count": len(teachers),
    })


#/ ─── Salary PDF generation / Зарплатная ведомость PDF ────────────────────────────────
#! Generates a downloadable PDF salary document with teacher list and totals
#! Генерирует PDF-ведомость со списком учителей и итогами
@router.get("/reports/salary/pdf", response_class=Response)
async def salary_pdf(
    request: Request,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(require_role("admin")),
):
    from reportlab.lib.units import mm, cm
    from reportlab.pdfgen import canvas
    from reportlab.lib import colors

    #* Fetch teachers
    result = await session.execute(
        select(User).where(User.role == "teacher")
    )
    teachers = result.scalars().all()
    teachers.sort(key=lambda u: (u.last_name or "", u.first_name or ""))

    buffer = io.BytesIO()
    c = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    margin = 2 * cm
    y = height - margin

    #* Try to find a Cyrillic-capable font
    font_name = "Helvetica"
    font_paths = [
        "/System/Library/Fonts/Supplemental/Arial.ttf",
        "/System/Library/Fonts/Supplemental/Verdana.ttf",
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
    ]
    for fp in font_paths:
        if os.path.exists(fp):
            from reportlab.pdfbase.ttfonts import TTFont
            from reportlab.pdfbase import pdfmetrics
            pdfmetrics.registerFont(TTFont("CyrillicFont", fp))
            font_name = "CyrillicFont"
            break

    def draw_header():
        nonlocal y
        c.setFont(font_name, 16)
        c.drawCentredString(width / 2, y, config.SCHOOL_NAME)
        y -= 20
        c.setFont(font_name, 12)
        c.drawCentredString(width / 2, y, "Зарплатная ведомость")
        y -= 16
        c.setFont(font_name, 9)
        today = date.today().strftime("%d.%m.%Y")
        c.drawCentredString(width / 2, y, f"Дата формирования: {today}")
        y -= 20

    def draw_table():
        nonlocal y
        col_widths = [20, 160, 60, 80, 60, 70]
        headers = ["№", "ФИО", "Стаж", "Зарплата\n(руб/мес)", "Часов\nв нед.", "Ставка\n(руб/час)"]
        row_h = 28
        table_top = y

        #* Check if we need a new page
        needed = len(teachers) * row_h + 120
        if y - needed < margin:
            c.showPage()
            y = height - margin
            draw_header()
            table_top = y

        #* Header row
        c.setFont(font_name, 8)
        c.setFillColor(colors.HexColor("#1f2937"))
        x = margin
        for i, (w, hdr) in enumerate(zip(col_widths, headers)):
            c.drawString(x, y - 10, hdr.replace("\n", " "))
            x += w
        y -= row_h

        #* Separator
        c.setStrokeColor(colors.HexColor("#d1d5db"))
        c.line(margin, y, width - margin, y)
        y -= 6

        total_salary = 0
        c.setFont(font_name, 8)
        for idx, t in enumerate(teachers, 1):
            c.setFillColor(colors.HexColor("#374151"))
            x = margin
            c.drawString(x, y - 10, str(idx))
            x += col_widths[0]
            name = f"{t.last_name} {t.first_name} {t.middle_name or ''}"
            c.drawString(x, y - 10, name[:30])

            x += col_widths[1]
            c.drawString(x, y - 10, str(t.experience_years or "—"))

            x += col_widths[2]
            salary = t.salary_monthly or 0
            total_salary += salary
            c.drawString(x, y - 10, f"{salary:,.0f}")

            x += col_widths[3]
            c.drawString(x, y - 10, str(t.hours_per_week or "—"))

            x += col_widths[4]
            if t.salary_monthly and t.hours_per_week and t.hours_per_week > 0:
                rate = t.salary_monthly / (t.hours_per_week * 4.33)
                c.drawString(x, y - 10, f"{rate:.2f}")
            else:
                c.drawString(x, y - 10, "—")

            y -= row_h

            #* Page break if needed
            if y < margin + 40:
                c.showPage()
                y = height - margin
                draw_header()

        #* Totals row
        y -= 4
        c.setStrokeColor(colors.HexColor("#374151"))
        c.line(margin, y, width - margin, y)
        y -= 16
        c.setFont(font_name, 9)
        c.setFillColor(colors.HexColor("#111827"))
        c.drawString(margin, y - 10, f"ИТОГО: {total_salary:,.0f} руб/мес")
        y -= 16
        if teachers:
            avg = total_salary / len(teachers)
            c.drawString(margin, y - 10, f"Средняя: {avg:,.0f} руб/мес")
        y -= 16
        c.drawString(margin, y - 10, f"Количество учителей: {len(teachers)}")

    draw_header()
    draw_table()

    #* Signature lines
    y -= 30
    c.setFont(font_name, 9)
    c.drawString(margin, y, "Директор: ___________________")
    c.drawString(width - margin - 120, y, "Бухгалтер: ___________________")
    y -= 20
    c.drawString(margin, y, "М.П.")
    c.drawString(width - margin - 140, y, f'Дата: "{today}"')

    c.save()
    buffer.seek(0)
    return Response(
        content=buffer.getvalue(),
        media_type="application/pdf",
        headers={
            "Content-Disposition": f'attachment; filename="salary_report_{date.today().strftime("%Y%m%d")}.pdf"'
        },
    )


#/ ═══════════════════════════════════════════════════════════════════════════════════════
#/  CLASS MANAGEMENT / УПРАВЛЕНИЕ КЛАССАМИ
#/ ═══════════════════════════════════════════════════════════════════════════════════════

#/ GET /admin/classes — List all classes with homeroom teacher
#/ Список всех классов с классными руководителями
@router.get("/classes", response_class=HTMLResponse)
async def list_classes(
    request: Request,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(require_role("director")),
):
    result = await session.execute(
        select(Class).options(selectinload(Class.homeroom_teacher)).order_by(Class.grade_level, Class.name)
    )
    classes = result.scalars().all()
    teachers = (await session.execute(select(User).where(User.role == "teacher"))).scalars().all()
    return templates.TemplateResponse("admin/classes.html", {
        "request": request, "user": user, "classes": classes, "teachers": teachers
    })


#! POST /admin/classes/create — Create new class + log
#! Создание нового класса + логирование
@router.post("/classes/create")
async def create_class(
    request: Request,
    session: AsyncSession = Depends(get_session),
    admin: User = Depends(require_role("director")),
):
    form = await request.form()
    name = form.get("name", "").strip()
    grade_level = int(form.get("grade_level", 1))
    homeroom_teacher_id = form.get("homeroom_teacher_id") or None
    academic_year = form.get("academic_year", "2024/2025")

    if not name:
        return RedirectResponse(url="/admin/classes", status_code=302)

    class_ = Class(name=name, grade_level=grade_level, homeroom_teacher_id=homeroom_teacher_id, academic_year=academic_year)
    session.add(class_)
    await session.commit()

    log_action(user_id=str(admin.id), action=CREATE_CLASS, details={"class_name": name})
    return RedirectResponse(url="/admin/classes", status_code=302)


#! POST /admin/classes/{class_id}/edit — Edit class details
#! Редактирование данных класса
@router.post("/classes/{class_id}/edit")
async def edit_class(
    request: Request, class_id: str,
    session: AsyncSession = Depends(get_session),
    admin: User = Depends(require_role("director")),
):
    result = await session.execute(select(Class).where(Class.id == class_id))
    class_ = result.scalar_one_or_none()
    if not class_:
        raise HTTPException(status_code=404)

    form = await request.form()
    class_.name = form.get("name", class_.name).strip()
    class_.grade_level = int(form.get("grade_level", class_.grade_level))
    class_.homeroom_teacher_id = form.get("homeroom_teacher_id") or None
    class_.academic_year = form.get("academic_year", class_.academic_year)
    await session.commit()
    return RedirectResponse(url="/admin/classes", status_code=302)


#! POST /admin/classes/{class_id}/delete — Delete class
#! Удаление класса
@router.post("/classes/{class_id}/delete")
async def delete_class(
    request: Request, class_id: str,
    session: AsyncSession = Depends(get_session),
    admin: User = Depends(require_role("director")),
):
    result = await session.execute(select(Class).where(Class.id == class_id))
    class_ = result.scalar_one_or_none()
    if class_:
        await session.delete(class_)
        await session.commit()
    return RedirectResponse(url="/admin/classes", status_code=302)


#/ ─── Auto-generate passwords for class / Автогенерация паролей для класса ────────────
#! Generates and hashes new passwords for all students, returns in table
#! Генерирует и хэширует новые пароли для всех учеников, выводит таблицу
@router.post("/classes/{class_id}/generate-credentials")
async def generate_class_credentials(
    request: Request, class_id: str,
    session: AsyncSession = Depends(get_session),
    admin: User = Depends(require_role("director")),
):
    #* Get the class and all students
    class_result = await session.execute(select(Class).where(Class.id == class_id))
    class_ = class_result.scalar_one_or_none()

    result = await session.execute(
        select(User).join(Enrollment).where(Enrollment.class_id == class_id, User.role == "student")
    )
    students = result.scalars().all()

    credentials = []
    for student in students:
        password = generate_password(8)
        student.password_hash = hash_password(password)
        await session.commit()

        credentials.append({
            "id": str(student.id),
            "first_name": student.first_name,
            "last_name": student.last_name,
            "middle_name": student.middle_name or "",
            "login": student.username or student.email,
            "password": password,
        })

    log_action(user_id=str(admin.id), action="generate_credentials",
               details={"class_id": class_id, "count": len(credentials)})

    return templates.TemplateResponse("admin/credentials_result.html", {
        "request": request, "user": admin,
        "credentials": credentials, "class_name": class_.name if class_ else ""
    })


#/ ─── Export credentials to DOCX / Экспорт паролей в DOCX ────────────────────────────
#! Generates DOCX with student logins (passwords hidden) using python-docx
#! Генерирует DOCX с логинами учеников (пароли скрыты) через python-docx
@router.post("/classes/{class_id}/export-docx")
async def export_credentials_docx(
    request: Request, class_id: str,
    session: AsyncSession = Depends(get_session),
    admin: User = Depends(require_role("director")),
):
    try:
        from docx import Document
        from docx.shared import Inches, Pt, Cm
        from docx.enum.table import WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        return HTMLResponse("Установите python-docx: pip install python-docx", status_code=500)

    #* Get class and students
    class_result = await session.execute(select(Class).where(Class.id == class_id))
    class_ = class_result.scalar_one_or_none()
    if not class_:
        raise HTTPException(status_code=404)

    result = await session.execute(
        select(User).join(Enrollment).where(Enrollment.class_id == class_id, User.role == "student")
    )
    students = result.scalars().all()

    doc = Document()

    #* Title
    title = doc.add_heading(f'Логины и пароли — {class_.name}', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #* Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr = table.rows[0].cells
    hdr[0].text = '№'
    hdr[1].text = 'ФИО'
    hdr[2].text = 'Логин'
    hdr[3].text = 'Пароль'

    for i, student in enumerate(students, 1):
        row = table.add_row().cells
        row[0].text = str(i)
        fio = f"{student.last_name} {student.first_name}"
        if student.middle_name:
            fio += f" {student.middle_name}"
        row[1].text = fio
        row[2].text = student.username or student.email
        row[3].text = "••••••••"  #* Don't show password in the document

    #* Column width settings
    for row in table.rows:
        row.cells[0].width = Cm(1)
        row.cells[1].width = Cm(6)
        row.cells[2].width = Cm(5)
        row.cells[3].width = Cm(3)

    #* Save to buffer
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)

    from starlette.responses import StreamingResponse
    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename=credentials_{class_.name}.docx"}
    )


#/ ═══════════════════════════════════════════════════════════════════════════════════════
#/  SUBJECT MANAGEMENT / УПРАВЛЕНИЕ ПРЕДМЕТАМИ
#/ ═══════════════════════════════════════════════════════════════════════════════════════

#/ GET /admin/subjects — List all subjects
#/ Список всех предметов
@router.get("/subjects", response_class=HTMLResponse)
async def list_subjects(
    request: Request,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(require_role("director")),
):
    result = await session.execute(select(Subject).order_by(Subject.name))
    subjects = result.scalars().all()
    return templates.TemplateResponse("admin/subjects.html", {
        "request": request, "user": user, "subjects": subjects
    })


#! POST /admin/subjects/create — Create new subject + log
#! Создание нового предмета + логирование
@router.post("/subjects/create")
async def create_subject(
    request: Request,
    session: AsyncSession = Depends(get_session),
    admin: User = Depends(require_role("director")),
):
    form = await request.form()
    name = form.get("name", "").strip()
    code = form.get("code", "").strip() or None
    if name:
        subject = Subject(name=name, code=code)
        session.add(subject)
        await session.commit()
        log_action(user_id=str(admin.id), action=CREATE_SUBJECT, details={"name": name})
    return RedirectResponse(url="/admin/subjects", status_code=302)


#! POST /admin/subjects/{subject_id}/delete — Delete subject
#! Удаление предмета
@router.post("/subjects/{subject_id}/delete")
async def delete_subject(
    request: Request, subject_id: str,
    session: AsyncSession = Depends(get_session),
    admin: User = Depends(require_role("director")),
):
    result = await session.execute(select(Subject).where(Subject.id == subject_id))
    subject = result.scalar_one_or_none()
    if subject:
        await session.delete(subject)
        await session.commit()
    return RedirectResponse(url="/admin/subjects", status_code=302)


#/ ═══════════════════════════════════════════════════════════════════════════════════════
#/  TEACHER ASSIGNMENTS / НАЗНАЧЕНИЕ УЧИТЕЛЕЙ
#/  Which teacher teaches which subject in which class / Какой учитель ведёт какой предмет в каком классе
#/ ═══════════════════════════════════════════════════════════════════════════════════════

#/ GET /admin/assignments — List all teacher-subject-class assignments
#/ Список всех назначений учитель-предмет-класс
@router.get("/assignments", response_class=HTMLResponse)
async def list_assignments(
    request: Request,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(require_role("director")),
):
    result = await session.execute(
        select(TeacherAssignment)
        .options(selectinload(TeacherAssignment.teacher), selectinload(TeacherAssignment.subject), selectinload(TeacherAssignment.class_))
    )
    assignments = result.scalars().all()
    teachers = (await session.execute(select(User).where(User.role == "teacher"))).scalars().all()
    subjects = (await session.execute(select(Subject))).scalars().all()
    classes = (await session.execute(select(Class))).scalars().all()

    return templates.TemplateResponse("admin/assignments.html", {
        "request": request, "user": user,
        "assignments": assignments, "teachers": teachers,
        "subjects": subjects, "classes": classes
    })


#! POST /admin/assignments/create — Assign teacher to subject+class
#! Назначение учителя на предмет+класс
@router.post("/assignments/create")
async def create_assignment(
    request: Request,
    session: AsyncSession = Depends(get_session),
    admin: User = Depends(require_role("director")),
):
    form = await request.form()
    teacher_id = form.get("teacher_id")
    subject_id = form.get("subject_id")
    class_id = form.get("class_id")

    if teacher_id and subject_id and class_id:
        assignment = TeacherAssignment(teacher_id=teacher_id, subject_id=subject_id, class_id=class_id)
        session.add(assignment)
        await session.commit()
        log_action(user_id=str(admin.id), action=ASSIGN_TEACHER,
                   details={"teacher_id": teacher_id, "subject_id": subject_id, "class_id": class_id})
    return RedirectResponse(url="/admin/assignments", status_code=302)


#! POST /admin/assignments/{assignment_id}/delete — Remove assignment
#! Удаление назначения
@router.post("/assignments/{assignment_id}/delete")
async def delete_assignment(
    request: Request, assignment_id: str,
    session: AsyncSession = Depends(get_session),
    admin: User = Depends(require_role("director")),
):
    result = await session.execute(select(TeacherAssignment).where(TeacherAssignment.id == assignment_id))
    assignment = result.scalar_one_or_none()
    if assignment:
        await session.delete(assignment)
        await session.commit()
    return RedirectResponse(url="/admin/assignments", status_code=302)


#/ ─── Schedule management / Управление расписанием ──────────────────────────────────────
#/ Schedule editor: add/remove lessons per class with conflict check
#/ Редактор расписания: добавление/удаление уроков с проверкой конфликтов
@router.get("/schedule", response_class=HTMLResponse)
async def manage_schedule(
    request: Request,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(require_role("secretary")),
):
    classes = (await session.execute(select(Class).order_by(Class.name))).scalars().all()
    subjects = (await session.execute(select(Subject).order_by(Subject.name))).scalars().all()
    teachers = (await session.execute(select(User).where(User.role == "teacher"))).scalars().all()
    teachers.sort(key=lambda u: u.last_name or "")

    #* Load schedule for selected class (only active entries)
    class_id = request.query_params.get("class_id")
    schedules = []
    if class_id:
        result = await session.execute(
            select(Schedule)
            .options(selectinload(Schedule.subject), selectinload(Schedule.teacher))
            .where(Schedule.class_id == class_id, Schedule.valid_until.is_(None))
            .order_by(Schedule.day_of_week, Schedule.lesson_number)
        )
        schedules = result.scalars().all()

    return templates.TemplateResponse("admin/schedule.html", {
        "request": request, "user": user,
        "classes": classes, "subjects": subjects, "teachers": teachers,
        "schedules": schedules, "selected_class_id": class_id
    })


#! POST /admin/schedule/add — Add lesson to schedule (with teacher conflict check)
#! Добавление урока в расписание (с проверкой конфликта учителя)
@router.post("/schedule/add")
async def add_schedule(
    request: Request,
    session: AsyncSession = Depends(get_session),
    secretary: User = Depends(require_role("secretary")),
):
    form = await request.form()
    class_id = form.get("class_id")
    teacher_id = form.get("teacher_id")
    day_of_week = int(form.get("day_of_week", 1))
    lesson_number = int(form.get("lesson_number", 1))

    #* Check teacher conflict against active schedules only
    conflict = await session.execute(
        select(Schedule).where(
            Schedule.teacher_id == teacher_id,
            Schedule.day_of_week == day_of_week,
            Schedule.lesson_number == lesson_number,
            Schedule.valid_until.is_(None),
        )
    )
    if conflict.scalar_one_or_none():
        classes = await session.execute(select(Class).order_by(Class.name))
        subjects = await session.execute(select(Subject).order_by(Subject.name))
        teachers = await session.execute(select(User).where(User.role == "teacher"))
        teachers_list = teachers.scalars().all()
        teachers_list.sort(key=lambda u: u.last_name or "")
        schedules = await session.execute(
            select(Schedule).where(Schedule.class_id == class_id, Schedule.valid_until.is_(None))
            .options(selectinload(Schedule.subject), selectinload(Schedule.teacher))
            .order_by(Schedule.day_of_week, Schedule.lesson_number)
        )
        return templates.TemplateResponse("admin/schedule.html", {
            "request": request, "user": secretary,
            "classes": classes.scalars().all(),
            "subjects": subjects.scalars().all(),
            "teachers": teachers_list,
            "schedules": schedules.scalars().all(),
            "selected_class_id": class_id,
            "error": "Учитель уже занят в это время в другом классе",
        }, status_code=400)

    schedule = Schedule(
        class_id=class_id,
        subject_id=form.get("subject_id"),
        teacher_id=teacher_id,
        day_of_week=day_of_week,
        lesson_number=lesson_number,
        classroom=form.get("classroom", "").strip() or None,
        valid_from=date.today(),
    )
    session.add(schedule)
    await session.commit()
    log_action(user_id=str(secretary.id), action="update_schedule",
               details={"class_id": str(schedule.class_id)})
    return RedirectResponse(url=f"/admin/schedule?class_id={schedule.class_id}", status_code=302)


#! POST /admin/schedule/{schedule_id}/delete — Soft-delete schedule entry
#! Мягкое удаление записи расписания (valid_until = yesterday)
@router.post("/schedule/{schedule_id}/delete")
async def delete_schedule(
    request: Request, schedule_id: str,
    session: AsyncSession = Depends(get_session),
    secretary: User = Depends(require_role("secretary")),
):
    result = await session.execute(select(Schedule).where(Schedule.id == schedule_id))
    sched = result.scalar_one_or_none()
    class_id = str(sched.class_id) if sched else ""
    if sched:
        sched.valid_until = date.today() - timedelta(days=1)
        await session.commit()
        log_action(user_id=str(secretary.id), action="delete_schedule")
    return RedirectResponse(url=f"/admin/schedule?class_id={class_id}", status_code=302)


#/ ═══════════════════════════════════════════════════════════════════════════════════════
#/  LOGS — view activity logs (director only) / ПРОСМОТР ЛОГОВ (только директор)
#/ ═══════════════════════════════════════════════════════════════════════════════════════

#! GET /admin/logs — View activity logs from MongoDB (admin/director only)
#! Просмотр логов активности из MongoDB (только администратор/директор)
@router.get("/logs", response_class=HTMLResponse)
async def view_logs(
    request: Request,
    user: User = Depends(require_auth),
):
    if user.role != "admin":
        raise HTTPException(status_code=403, detail="Только администратор может просматривать логи")
    #* Get logs from MongoDB
    logs = []
    try:
        from app.database import get_mongo
        mongo = get_mongo()
        if mongo:
            collection = mongo["school_logs"]["activity_logs"]
            cursor = collection.find().sort("timestamp", -1).limit(500)
            logs = list(cursor)
            #* Convert ObjectId to string and format time
            for log in logs:
                log["_id"] = str(log["_id"])
                if "timestamp" in log:
                    log["timestamp"] = log["timestamp"].strftime("%d.%m.%Y %H:%M")
    except Exception as e:
        print(f"[!] Failed to fetch logs: {e}")

    log_action(user_id=str(user.id), action=VIEW_LOGS)
    return templates.TemplateResponse("admin/logs.html", {
        "request": request, "user": user, "logs": logs
    })


#/ ─── Manual log cleanup / Ручная очистка логов ──────────────────────────────────────
#! Deletes all logs older than 365 days, then re-logs the cleanup action
#! Удаляет все логи старше 365 дней и логирует действие очистки
@router.post("/logs/cleanup")
async def cleanup_logs(
    request: Request,
    user: User = Depends(require_role("admin")),
):
    try:
        from app.database import get_mongo
        mongo = get_mongo()
        count = 0
        if mongo:
            from datetime import timedelta
            cleanup_before = (datetime.now(timezone.utc) - timedelta(days=365)).replace(
                hour=0, minute=0, second=0, microsecond=0
            )
            result = mongo["school_logs"]["activity_logs"].delete_many(
                {"timestamp": {"$lt": cleanup_before}}
            )
            count = result.deleted_count
        log_action(user_id=str(user.id), action=CLEANUP_LOGS,
                   details={"deleted_count": count})
        return RedirectResponse(url="/admin/logs?cleaned=1", status_code=302)
    except Exception as e:
        return HTMLResponse(f"Ошибка очистки: {e}", status_code=500)


#/ ═══════════════════════════════════════════════════════════════════════════════════════
#/  TEACHERS — list and edit / УЧИТЕЛЯ — список и редактирование
#/ ═══════════════════════════════════════════════════════════════════════════════════════

#/ GET /admin/teachers — List teachers with their assignments
#/ Список учителей с их назначениями
@router.get("/teachers", response_class=HTMLResponse)
async def list_teachers(
    request: Request,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(require_role("director")),
):
    result = await session.execute(
        select(User).where(User.role == "teacher")
    )
    teachers = result.scalars().all()
    teachers.sort(key=lambda u: (u.last_name or "", u.first_name or ""))

    #* Get assignments for each teacher
    teacher_assignments = {}
    for t in teachers:
        assignments = await session.execute(
            select(TeacherAssignment)
            .options(selectinload(TeacherAssignment.subject), selectinload(TeacherAssignment.class_))
            .where(TeacherAssignment.teacher_id == t.id)
        )
        teacher_assignments[str(t.id)] = assignments.scalars().all()

    return templates.TemplateResponse("admin/teachers.html", {
        "request": request, "user": user,
        "teachers": teachers, "teacher_assignments": teacher_assignments,
    })


#/ GET /admin/teachers/{teacher_id}/edit — Edit teacher form
#/ Форма редактирования учителя
@router.get("/teachers/{teacher_id}/edit", response_class=HTMLResponse)
async def edit_teacher_page(
    request: Request, teacher_id: str,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(require_role("director")),
):
    result = await session.execute(select(User).where(User.id == teacher_id))
    target = result.scalar_one_or_none()
    if not target or target.role != "teacher":
        raise HTTPException(status_code=404, detail="Учитель не найден")
    return templates.TemplateResponse("admin/teacher_edit.html", {
        "request": request, "user": user, "target": target
    })


#! POST /admin/teachers/{teacher_id}/edit — Update teacher data + log
#! Обновление данных учителя + логирование
@router.post("/teachers/{teacher_id}/edit")
async def edit_teacher(
    request: Request, teacher_id: str,
    session: AsyncSession = Depends(get_session),
    user: User = Depends(require_role("admin")),
):
    result = await session.execute(select(User).where(User.id == teacher_id))
    target = result.scalar_one_or_none()
    if not target or target.role != "teacher":
        raise HTTPException(status_code=404)

    form = await request.form()
    target.first_name = form.get("first_name", target.first_name).strip()
    target.last_name = form.get("last_name", target.last_name).strip()
    target.middle_name = form.get("middle_name", "").strip() or None
    target.email = form.get("email", target.email).strip().lower()
    target.phone = form.get("phone", "").strip() or None
    target.labor_book_number = form.get("labor_book_number", "").strip() or None
    target.experience_years = int(form.get("experience_years", 0)) if form.get("experience_years", "").strip() else None
    salary_str = form.get("salary_monthly", "").strip()
    target.salary_monthly = float(salary_str) if salary_str else None
    hours_str = form.get("hours_per_week", "").strip()
    target.hours_per_week = float(hours_str) if hours_str else None
    target.staff_position = form.get("staff_position", "").strip() or None
    target.address = form.get("address", "").strip() or None

    new_password = form.get("password", "").strip()
    if new_password:
        target.password_hash = hash_password(new_password)

    await session.commit()
    log_action(
        user_id=str(user.id), action=UPDATE_USER,
        details={
            "target_user_id": teacher_id,
            "fields": {
                "first_name": target.first_name,
                "last_name": target.last_name,
                "salary_monthly": target.salary_monthly,
                "hours_per_week": target.hours_per_week,
                "experience_years": target.experience_years,
                "labor_book_number": target.labor_book_number,
            }
        },
    )
    return RedirectResponse(url="/admin/reports/salary", status_code=302)
